# src/extractors/text_extractor.py

from typing import List, Dict, Any, Tuple
import pandas as pd
import numpy as np
from pathlib import Path
from pypdf import PdfReader
import docx
import re
import logging
from datetime import datetime

class DocumentProcessor:
    def __init__(self, config: Dict[str, Any]):
        """初始化文檔處理器"""
        self.chunk_size = config.get('chunk_size', 1000)
        self.chunk_overlap = config.get('chunk_overlap', 200)
        self.logger = logging.getLogger(__name__)
        
    def process_document(self, file_path: Path) -> Tuple[List[str], Dict[str, Any]]:
        """
        處理文檔並返回文本塊和元數據
        
        Args:
            file_path: 文檔路徑
            
        Returns:
            (文本塊列表, 元數據字典)
        """
        try:
            # 根據文件類型選擇處理方法
            suffix = file_path.suffix.lower()
            
            # 初始化元數據
            metadata = {
                'file_name': file_path.name,
                'file_type': suffix,
                'process_time': datetime.now().isoformat(),
                'tables_found': 0,
                'total_chunks': 0
            }
            
            if suffix == '.pdf':
                text, tables = self._read_pdf(file_path)
            elif suffix == '.docx':
                text, tables = self._read_docx(file_path)
            elif suffix == '.xlsx':
                text, tables = self._read_excel(file_path)
            else:
                raise ValueError(f"Unsupported file type: {suffix}")
                
            # 提取表格中的文本
            table_texts = self._process_tables(tables)
            
            # 合併所有文本並分割
            all_text = text + "\n" + table_texts
            chunks = self._split_text(all_text)
            
            # 更新元數據
            metadata['tables_found'] = len(tables)
            metadata['total_chunks'] = len(chunks)
            
            return chunks, metadata
            
        except Exception as e:
            self.logger.error(f"Error processing document: {str(e)}")
            raise
            
    def _read_pdf(self, file_path: Path) -> Tuple[str, List[pd.DataFrame]]:
        """
        讀取 PDF 文件
        
        Returns:
            (文本內容, 表格列表)
        """
        reader = PdfReader(str(file_path))
        text = ""
        tables = []
        
        for page in reader.pages:
            # 提取文本
            text += page.extract_text() + "\n"
            
            # 嘗試提取表格
            try:
                # 使用 tabula-py 提取表格
                page_tables = self._extract_tables_from_pdf_page(page)
                tables.extend(page_tables)
            except Exception as e:
                self.logger.warning(f"Error extracting tables from page: {str(e)}")
                
        return text, tables
        
    def _read_docx(self, file_path: Path) -> Tuple[str, List[pd.DataFrame]]:
        """
        讀取 Word 文件
        
        Returns:
            (文本內容, 表格列表)
        """
        doc = docx.Document(str(file_path))
        text = ""
        tables = []
        
        # 處理段落
        for para in doc.paragraphs:
            text += para.text + "\n"
            
        # 處理表格
        for table in doc.tables:
            try:
                # 將 Word 表格轉換為 DataFrame
                data = []
                for row in table.rows:
                    data.append([cell.text for cell in row.cells])
                if data:
                    df = pd.DataFrame(data[1:], columns=data[0])
                    tables.append(df)
            except Exception as e:
                self.logger.warning(f"Error processing table: {str(e)}")
                
        return text, tables
        
    def _read_excel(self, file_path: Path) -> Tuple[str, List[pd.DataFrame]]:
        """
        讀取 Excel 文件
        
        Returns:
            (文本內容, 表格列表)
        """
        # 讀取所有工作表
        excel_file = pd.ExcelFile(str(file_path))
        all_sheets = []
        text = ""
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet_name)
            all_sheets.append(df)
            
            # 將表格轉換為文本描述
            text += f"\n工作表: {sheet_name}\n"
            text += df.to_string(index=False) + "\n"
            
        return text, all_sheets
        
    def _process_tables(self, tables: List[pd.DataFrame]) -> str:
        """
        處理表格並生成結構化文本描述
        """
        text = ""
        for i, df in enumerate(tables, 1):
            text += f"\n表格 {i}:\n"
            # 添加列名描述
            text += f"欄位: {', '.join(df.columns)}\n"
            # 添加數據描述
            text += df.to_string(index=False) + "\n"
            
        return text
        
    def _split_text(self, text: str) -> List[str]:
        """
        將文本分割成塊
        """
        chunks = []
        current_chunk = ""
        
        # 使用正則表達式匹配可能的章節標題
        section_pattern = r'\n(?:[一二三四五六七八九十]+、|\d+\.|第[一二三四五六七八九十]+章)'
        
        # 按行分割文本
        lines = text.split('\n')
        
        for line in lines:
            # 如果遇到新章節且當前塊不為空
            if re.match(section_pattern, line) and current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = line
            else:
                # 如果當前塊長度即將超過限制
                if len(current_chunk) + len(line) > self.chunk_size - self.chunk_overlap:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = line
                else:
                    current_chunk += "\n" + line
                    
        # 添加最後一個塊
        if current_chunk:
            chunks.append(current_chunk.strip())
            
        return chunks

    def _extract_metric_candidates(self, text: str) -> List[Dict[str, Any]]:
        """
        從文本中提取可能的指標候選項
        """
        candidates = []
        
        # 匹配數字模式（包括百分比）
        number_pattern = r'(\d+(?:\.\d+)?)\s*(?:%|％|萬|億|個|人|元|小時|分鐘)?'
        
        # 匹配常見的 ESG 相關詞彙
        esg_keywords = ['員工', '環境', '碳排', '能源', '培訓', '薪資', '安全', '事故']
        
        for keyword in esg_keywords:
            # 在關鍵詞周圍尋找數字
            matches = re.finditer(fr'{keyword}[^。\n]*?{number_pattern}', text)
            for match in matches:
                candidates.append({
                    'keyword': keyword,
                    'context': match.group(0),
                    'value': match.group(1),
                    'unit': match.group(2) if len(match.groups()) > 1 else None
                })
                
        return candidates